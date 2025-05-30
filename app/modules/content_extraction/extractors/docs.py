from docx import Document
from typing import Dict
from pathlib import Path

from app.core.logging_config import get_logger

logger = get_logger(__name__)

def extract_docx_text(input_path: str, output_path: str) -> Dict:
    """Extract text from DOCX files"""
    logger.info(f"Extracting docx file [start]")
    try:
        doc = Document(input_path)
        text = "\n".join([para.text for para in doc.paragraphs])

        # # Extract and print the text from each paragraph
        # for para in doc.paragraphs:
        #     print(para.text)
        #     logger.info("Extracted text:", para.text)

        # logger.info(f"Extracted text: {text}")
        
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
        
        logger.info(f"Extracted [saved]: {output_path}")

        return {
            "content_type": "text",
            "method": "python-docx",
            "word_count": len(text.split())
        }
    except Exception as e:
        logger.error(f"Failed to extract content from docx: {str(e)}", exc_info=True)
        raise Exception(f"DOCX extraction failed: {str(e)}")
    

if __name__ == "__main__":

    path = r"C:\Users\ask50\Desktop\MP Report\Certi.docx"
    output = "D:\PROJECT\DMS\Certi_output.txt"
    extract_docx_text(path, output)
    print("Extracted")